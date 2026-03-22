import io
import os
import csv
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from fastapi import UploadFile

SUPPORTED_FILE_KINDS: tuple[str, ...] = (
    "csv",
    "xls",
    "xlsx",
    "pdf",
    "docx",
    "png",
    "jpg",
    "tiff",
    "txt",
    "md",
    "rtf",
    "odt",
    "xml",
    "epub",
    "fb2",
    "doc",
    "unknown",
)


class ParseFileError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message

    def as_detail(self) -> Dict[str, str]:
        return {"code": self.code, "message": self.message}


def _unified_extracted_payload(
    *,
    file_kind: str,
    text: str = "",
    tables: Optional[List[Any]] = None,
    records: Optional[List[Dict[str, Any]]] = None,
    warnings: Optional[List[str]] = None,
    extra_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    metadata: Dict[str, Any] = {
        "kind": file_kind,
        "warnings": warnings or [],
        "has_text": bool((text or "").strip()),
        "records_count": len(records or []),
        "tables_count": len(tables or []),
    }
    if extra_metadata:
        metadata.update(extra_metadata)
    return {
        "kind": file_kind,
        "text": text or "",
        "tables": tables or [],
        "records": records or [],
        "metadata": metadata,
    }


def detect_file_kind(filename: Optional[str], content_type: Optional[str]) -> str:
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    # Prefer extension.
    if name.endswith(".csv") or "text/csv" in ctype:
        return "csv"
    if name.endswith(".xls") or "application/vnd.ms-excel" in ctype:
        return "xls"
    if name.endswith(".xlsx") or "excel" in ctype and "spreadsheetml" in ctype:
        return "xlsx"
    if name.endswith(".pdf") or "application/pdf" in ctype:
        return "pdf"
    if name.endswith(".docx") or "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in ctype:
        return "docx"
    if name.endswith(".png") or "image/png" in ctype:
        return "png"
    if name.endswith(".jpg") or name.endswith(".jpeg") or "image/jpeg" in ctype:
        return "jpg"
    if name.endswith(".tif") or name.endswith(".tiff") or "image/tiff" in ctype:
        return "tiff"
    if name.endswith(".txt") or "text/plain" in ctype:
        return "txt"
    if name.endswith(".md") or "text/markdown" in ctype or "text/x-markdown" in ctype:
        return "md"
    if name.endswith(".rtf") or "application/rtf" in ctype or "text/rtf" in ctype:
        return "rtf"
    if name.endswith(".odt") or "application/vnd.oasis.opendocument.text" in ctype:
        return "odt"
    if name.endswith(".xml") or "application/xml" in ctype or "text/xml" in ctype:
        return "xml"
    if name.endswith(".epub") or "application/epub+zip" in ctype:
        return "epub"
    if name.endswith(".fb2") or "application/x-fictionbook+xml" in ctype:
        return "fb2"
    if name.endswith(".doc") or "application/msword" in ctype:
        return "doc"

    return "unknown"


def _limit_records(records: List[Dict[str, Any]], max_rows: int) -> List[Dict[str, Any]]:
    if len(records) <= max_rows:
        return records
    return records[:max_rows]


def _merge_docx_kv_cells(prev: str, nxt: str) -> str:
    p = (prev or "").strip()
    n = (nxt or "").strip()
    if not p:
        return n
    if not n:
        return p
    if n in p:
        return p
    return f"{p}\n{n}"


def _disambiguate_docx_headers(headers: List[str]) -> List[str]:
    out: List[str] = []
    counts: Dict[str, int] = {}
    for i, h in enumerate(headers):
        base = (str(h or "").strip()) or f"col_{i + 1}"
        n = counts.get(base, 0)
        counts[base] = n + 1
        out.append(base if n == 0 else f"{base}::{n}")
    return out


def _docx_header_row_is_uniform(headers: List[str]) -> bool:
    cells = [str(h or "").strip() for h in headers]
    non_empty = [h for h in cells if h]
    return len(cells) >= 2 and len(set(non_empty)) <= 1 and bool(non_empty)


def _extract_kv_records_from_raw_rows(raw_rows: List[Any]) -> List[Dict[str, str]]:
    records: List[Dict[str, str]] = []
    kv: Dict[str, str] = {}
    current_key = ""

    def flush() -> None:
        nonlocal kv, current_key
        if kv:
            records.append(dict(kv))
        kv = {}
        current_key = ""

    for row in raw_rows:
        if not isinstance(row, list):
            continue
        cells = [str(c or "").strip() for c in row]
        if not any(cells):
            continue

        left = cells[0] if len(cells) >= 1 else ""
        right = cells[1] if len(cells) >= 2 else ""

        if left and right:
            if left in kv and str(kv.get(left, "")).strip() and str(right).strip():
                flush()
            current_key = left
            kv[left] = _merge_docx_kv_cells(kv.get(left, ""), right)
            continue

        if not left and right and current_key:
            kv[current_key] = _merge_docx_kv_cells(kv.get(current_key, ""), right)
            continue

        if current_key and len([c for c in cells if c]) == 1:
            val = next((c for c in cells if c), "")
            kv[current_key] = _merge_docx_kv_cells(kv.get(current_key, ""), val)

    if kv:
        records.append(kv)
    return records


def _to_records_dataframe(df: pd.DataFrame, max_rows: Optional[int]) -> List[Dict[str, Any]]:
    if max_rows is not None:
        df = df.head(max_rows)
    df = df.fillna("")
    records = df.to_dict(orient="records")

    for r in records:
        for k, v in list(r.items()):
            if v is None:
                r[k] = ""
            elif isinstance(v, (int, float, bool)):
                r[k] = str(v)
            else:
                r[k] = str(v)
    return records


def _detect_csv_delimiter(text_sample: str) -> str:
    """
    Detect delimiter from sample text, fallback to semicolon for CRM-style exports.
    """
    sample = (text_sample or "").strip()
    if not sample:
        return ";"
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        if getattr(dialect, "delimiter", None):
            return dialect.delimiter
    except Exception:
        pass
    return ";" if sample.count(";") >= sample.count(",") else ","


def _read_csv_dataframe(contents: bytes, *, max_rows: Optional[int]) -> Tuple[pd.DataFrame, str]:
    """
    Read CSV with delimiter/encoding tolerance.
    Returns (dataframe, detected_delimiter) for downstream TS generation.
    """

    decoded_text = ""
    for enc in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            decoded_text = contents.decode(enc)
            break
        except Exception:
            continue

    if not decoded_text:
        decoded_text = contents.decode("utf-8", errors="replace")

    delimiter = _detect_csv_delimiter(decoded_text[:8000])
    df = pd.read_csv(
        io.StringIO(decoded_text),
        sep=delimiter,
        nrows=max_rows,
        dtype=str,
        keep_default_na=False,
        engine="python",
    )
    return df, delimiter


def _decode_text_contents(contents: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return contents.decode(enc)
        except Exception:
            continue
    return contents.decode("utf-8", errors="replace")


def _normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _truncate_text(text: str, max_text_chars: Optional[int]) -> str:
    if max_text_chars is not None and max_text_chars > 0 and len(text) > max_text_chars:
        return text[: max_text_chars - 1] + "…"
    return text


def _extract_image_transcript(contents: bytes, file_kind: str) -> str:
    from app.services.image_transcription import transcribe_image_with_ocr

    try:
        text = transcribe_image_with_ocr(contents, file_kind).strip()
    except ValueError as e:
        raise ParseFileError(code="OCR_NO_TEXT", message=str(e)) from e
    except Exception as e:
        raise ParseFileError(
            code="OCR_FAILED",
            message=f"Не удалось выполнить OCR для изображения: {e}",
        ) from e

    if not text:
        raise ParseFileError(
            code="OCR_NO_TEXT",
            message="OCR не вернул текст с изображения. Проверьте качество снимка, язык OCR и контрастность.",
        )
    return text


def _join_non_empty_lines(lines: List[str]) -> str:
    return "\n".join([ln.strip() for ln in lines if ln and ln.strip()])


def _extract_rtf_text(contents: bytes) -> str:
    from striprtf.striprtf import rtf_to_text

    raw = _decode_text_contents(contents)
    return rtf_to_text(raw) or ""


def _extract_odt_text(contents: bytes) -> str:
    from odf import text as odf_text
    from odf.opendocument import load
    from odf.teletype import extractText

    doc = load(io.BytesIO(contents))
    lines: List[str] = []
    for elem in doc.getElementsByType(odf_text.P):
        val = extractText(elem)
        if val:
            lines.append(val)
    for elem in doc.getElementsByType(odf_text.H):
        val = extractText(elem)
        if val:
            lines.append(val)
    return _join_non_empty_lines(lines)


def _extract_xml_text(contents: bytes) -> str:
    xml_text = _decode_text_contents(contents)
    root = ET.fromstring(xml_text)
    pieces = [t.strip() for t in root.itertext() if t and t.strip()]
    return _join_non_empty_lines(pieces)


def _extract_epub_text(contents: bytes) -> str:
    from bs4 import BeautifulSoup
    from ebooklib import ITEM_DOCUMENT, epub

    with zipfile.ZipFile(io.BytesIO(contents), "r") as zf:
        with zf.open("META-INF/container.xml") as f:
            container_xml = f.read().decode("utf-8", errors="replace")
    container_root = ET.fromstring(container_xml)
    rootfile_path = ""
    for elem in container_root.iter():
        if elem.tag.lower().endswith("rootfile"):
            rootfile_path = (elem.attrib.get("full-path") or "").strip()
            break
    if not rootfile_path:
        raise ParseFileError(code="TEXT_DECODE_FAILED", message="EPUB container.xml has no rootfile path.")

    book = epub.read_epub(io.BytesIO(contents), options={"ignore_ncx": True})
    lines: List[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        html = item.get_content().decode("utf-8", errors="replace")
        text = BeautifulSoup(html, "html.parser").get_text(separator="\n", strip=True)
        if text:
            lines.append(text)
    return _join_non_empty_lines(lines)


def _extract_doc_text(contents: bytes) -> str:
    utf16 = contents.decode("utf-16le", errors="ignore")
    cp = contents.decode("cp1251", errors="ignore")
    mix = f"{utf16}\n{cp}"
    words = re.findall(r"[A-Za-zА-Яа-яЁё0-9][A-Za-zА-Яа-яЁё0-9 .,;:!?()\"'/-]{2,}", mix)
    return _join_non_empty_lines(words[:2000])


def _normalize_broken_semicolon_rows(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not records:
        return records

    normalized: List[Dict[str, Any]] = []
    for row in records:
        if not isinstance(row, dict) or len(row) != 1:
            normalized.append(row)
            continue

        only_key = next(iter(row.keys()), "")
        only_val = str(next(iter(row.values()), ""))
        if ";" not in str(only_key):
            normalized.append(row)
            continue

        headers = [h.strip() for h in str(only_key).split(";")]
        values = [v.strip() for v in only_val.split(";")]
        expanded: Dict[str, Any] = {}
        for i, h in enumerate(headers):
            if not h:
                continue
            expanded[h] = values[i] if i < len(values) else ""
        normalized.append(expanded if expanded else row)

    return normalized


def _records_from_text_key_value(text: str, *, max_rows: Optional[int]) -> List[Dict[str, Any]]:
    pairs: Dict[str, str] = {}
    for raw_line in _normalize_newlines(text).split("\n"):
        line = (raw_line or "").strip()
        if not line:
            continue
        m = re.match(r"^([^:;]{2,120})\s*[:\-]\s*(.+)$", line)
        if not m:
            continue
        key = m.group(1).strip()
        value = m.group(2).strip()
        if key and value:
            pairs[key] = value
    if not pairs:
        return []
    records = [pairs]
    return _limit_records(records, max_rows) if max_rows is not None else records


def _records_from_ocr_text(text: str, *, max_rows: Optional[int]) -> List[Dict[str, Any]]:
    kv_records = _records_from_text_key_value(text, max_rows=max_rows)
    if kv_records:
        return kv_records

    lines = [ln.strip() for ln in _normalize_newlines(text).split("\n") if ln and ln.strip()]
    if not lines:
        return []

    q_re = re.compile(r"^\d+[.)]\s*(.+)$")
    opt_re = re.compile(r"^[A-Za-zА-Яа-яЁё0-9][.)]?\s+(.+)$")
    records: List[Dict[str, Any]] = []
    current_question = ""
    current_options: List[str] = []

    def flush() -> None:
        nonlocal current_question, current_options
        if current_question:
            item: Dict[str, Any] = {"question": current_question}
            if current_options:
                item["options"] = " | ".join(current_options)
            records.append(item)
        current_question = ""
        current_options = []

    for ln in lines:
        q_m = q_re.match(ln)
        if q_m:
            flush()
            current_question = q_m.group(1).strip()
            continue
        o_m = opt_re.match(ln)
        if o_m and current_question:
            current_options.append(o_m.group(1).strip())
            continue
        if current_question:
            current_question = f"{current_question} {ln}".strip()

    flush()
    if records:
        return _limit_records(records, max_rows) if max_rows is not None else records

    text_preview = " ".join(lines[:6]).strip()
    fallback = [{"content": text_preview}] if text_preview else []
    return _limit_records(fallback, max_rows) if max_rows is not None else fallback


def _records_from_doc_tables(tables: List[Dict[str, Any]], *, max_rows: Optional[int]) -> List[Dict[str, Any]]:
    def _norm(s: Any) -> str:
        return re.sub(r"\s+", " ", str(s or "").strip().lower())

    def _is_meaningful_row_dict(row: Dict[str, Any]) -> bool:
        if not row:
            return False
        keys = [str(k or "").strip() for k in row.keys() if str(k or "").strip()]
        unique_keys = {_norm(k) for k in keys if _norm(k)}
        non_empty_values = [str(v or "").strip() for v in row.values() if str(v or "").strip()]

        if len(unique_keys) <= 1 and len(non_empty_values) <= 1:
            return False
        return True

    out: List[Dict[str, Any]] = []
    for table in tables:
        raw = table.get("raw") if isinstance(table, dict) else None
        if not isinstance(raw, list) or not raw:
            continue

        headers_raw = table.get("headers") if isinstance(table.get("headers"), list) else []
        hdr_cells = [str(h or "").strip() for h in headers_raw]
        uniform_hdr = _docx_header_row_is_uniform(hdr_cells)

        row_records: List[Dict[str, Any]] = []
        rows = table.get("rows") if isinstance(table, dict) else None
        if not uniform_hdr and isinstance(rows, list):
            for r in rows:
                if isinstance(r, dict) and _is_meaningful_row_dict(r):
                    row_records.append({str(k): str(v) for k, v in r.items()})

        raw_for_kv = raw[1:] if uniform_hdr and len(raw) > 1 else raw
        multi_kv = _extract_kv_records_from_raw_rows(raw_for_kv)

        if uniform_hdr:
            for kv in multi_kv:
                if kv:
                    out.append(kv)
        elif row_records:
            out.extend(row_records)
        else:
            for kv in multi_kv:
                if kv:
                    out.append(kv)

    if not out:
        return []
    return _limit_records(out, max_rows) if max_rows is not None else out


def extract_extracted_input_from_bytes(
    filename: Optional[str],
    content_type: Optional[str],
    contents: bytes,
    *,
    max_rows: Optional[int] = None,
    max_text_chars: Optional[int] = None,
) -> Tuple[str, Any]:

    file_kind = detect_file_kind(filename, content_type)

    if file_kind in {"csv", "xls", "xlsx"}:
        extra_meta: Dict[str, Any] = {}
        if file_kind == "csv":
            try:
                from app.services.sber_extract import parse_csv_dict_rows

                raw_rows, csv_delim = parse_csv_dict_rows(contents)
                extra_meta["csv_delimiter"] = csv_delim
                capped = raw_rows if max_rows is None else raw_rows[:max_rows]
                records = [{str(k): str(v) for k, v in r.items()} for r in capped]
                records = _normalize_broken_semicolon_rows(records)
                return file_kind, _unified_extracted_payload(
                    file_kind=file_kind,
                    records=records,
                    extra_metadata=extra_meta,
                )
            except Exception:
                df, csv_delim = _read_csv_dataframe(contents, max_rows=max_rows)
                extra_meta["csv_delimiter"] = csv_delim
                records = _to_records_dataframe(df, max_rows=max_rows)
                records = _normalize_broken_semicolon_rows(records)
                return file_kind, _unified_extracted_payload(
                    file_kind=file_kind,
                    records=records,
                    extra_metadata=extra_meta,
                )
        elif file_kind == "xlsx":
            try:
                from app.services.sber_extract import parse_xlsx_dict_rows

                raw_rows = parse_xlsx_dict_rows(contents, max_rows=max_rows)
                records = [{str(k): str(v) for k, v in r.items()} for r in raw_rows]
                records = _normalize_broken_semicolon_rows(records)
                return file_kind, _unified_extracted_payload(file_kind=file_kind, records=records)
            except Exception:
                bytes_buf = io.BytesIO(contents)
                df = pd.read_excel(bytes_buf, nrows=max_rows, dtype=str, sheet_name=0, engine=None)
                records = _to_records_dataframe(df, max_rows=max_rows)
                records = _normalize_broken_semicolon_rows(records)
                return file_kind, _unified_extracted_payload(file_kind=file_kind, records=records)
        elif file_kind == "xls":
            bytes_buf = io.BytesIO(contents)
            df = pd.read_excel(bytes_buf, nrows=max_rows, dtype=str, sheet_name=0, engine=None)
            records = _to_records_dataframe(df, max_rows=max_rows)
            records = _normalize_broken_semicolon_rows(records)
            return file_kind, _unified_extracted_payload(file_kind=file_kind, records=records)
        else:
            raise ValueError("Unsupported spreadsheet kind")

    if file_kind in {"txt", "md"}:
        text = _normalize_newlines(_decode_text_contents(contents))
        text = _truncate_text(text, max_text_chars)
        return file_kind, _unified_extracted_payload(file_kind=file_kind, text=text)

    if file_kind == "rtf":
        text = _normalize_newlines(_extract_rtf_text(contents))
        text = _truncate_text(text, max_text_chars)
        if not text.strip():
            raise ParseFileError(code="TEXT_DECODE_FAILED", message="Failed to extract text from RTF file.")
        return file_kind, _unified_extracted_payload(file_kind=file_kind, text=text)

    if file_kind == "odt":
        text = _normalize_newlines(_extract_odt_text(contents))
        text = _truncate_text(text, max_text_chars)
        if not text.strip():
            raise ParseFileError(code="TEXT_DECODE_FAILED", message="Failed to extract text from ODT file.")
        return file_kind, _unified_extracted_payload(file_kind=file_kind, text=text)

    if file_kind in {"xml", "fb2"}:
        text = _normalize_newlines(_extract_xml_text(contents))
        text = _truncate_text(text, max_text_chars)
        if not text.strip():
            raise ParseFileError(code="TEXT_DECODE_FAILED", message=f"Failed to extract text from {file_kind.upper()} file.")
        return file_kind, _unified_extracted_payload(file_kind=file_kind, text=text)

    if file_kind == "epub":
        text = _normalize_newlines(_extract_epub_text(contents))
        text = _truncate_text(text, max_text_chars)
        if not text.strip():
            raise ParseFileError(code="TEXT_DECODE_FAILED", message="Failed to extract text from EPUB file.")
        return file_kind, _unified_extracted_payload(file_kind=file_kind, text=text)

    if file_kind == "doc":
        text = _normalize_newlines(_extract_doc_text(contents))
        text = _truncate_text(text, max_text_chars)
        if not text.strip():
            raise ParseFileError(code="TEXT_DECODE_FAILED", message="Failed to extract text from DOC file.")
        return file_kind, _unified_extracted_payload(
            file_kind=file_kind,
            text=text,
            warnings=["legacy_doc_best_effort_extraction"],
        )

    if file_kind == "pdf":
        from app.services.sber_extract import extract_fatca_row_from_text, extract_pdf_text_pypdf, looks_like_fatca_text

        text = ""
        try:
            text = extract_pdf_text_pypdf(contents)
        except Exception:
            text = ""
        if not (text or "").strip():
            from PyPDF2 import PdfReader

            bytes_buf = io.BytesIO(contents)
            reader = PdfReader(bytes_buf)
            texts: List[str] = []
            pages = reader.pages if max_rows is None else reader.pages[: max(1, max_rows)]
            for page in pages:
                try:
                    texts.append(page.extract_text() or "")
                except Exception:
                    texts.append("")
            text = "\n".join(texts).strip()
        text = _truncate_text(_normalize_newlines(text), max_text_chars)
        if not text.strip():
            raise ParseFileError(
                code="TEXT_DECODE_FAILED",
                message="Failed to extract text from PDF file. The document may be scanned/image-only or protected.",
            )
        records = _records_from_text_key_value(text, max_rows=max_rows)
        if looks_like_fatca_text(text):
            fr = extract_fatca_row_from_text(text)
            if any(str(fr.get(k, "")).strip() for k in fr):
                records = [fr] + [r for r in (records or []) if r]
        return file_kind, _unified_extracted_payload(file_kind=file_kind, text=text, records=records)

    if file_kind == "docx":
        from docx import Document

        bytes_buf = io.BytesIO(contents)
        doc = Document(bytes_buf)
        paragraphs: List[str] = []
        for p in doc.paragraphs:
            if p.text:
                paragraphs.append(p.text)
        tables: List[Dict[str, Any]] = []
        tables_iter = doc.tables if max_rows is None else doc.tables[: max(1, max_rows)]
        for t in tables_iter:
            table_data: List[List[str]] = []
            for row in t.rows:
                row_data: List[str] = []
                for cell in row.cells:
                    row_data.append((cell.text or "").strip())
                table_data.append(row_data)
            headers_raw: List[str] = [str(c or "") for c in (table_data[0] if table_data else [])]
            headers = _disambiguate_docx_headers(headers_raw)
            rows: List[Dict[str, str]] = []
            if headers and len(table_data) > 1 and any(str(h or "").strip() for h in headers):
                for raw_row in table_data[1:]:
                    row_obj: Dict[str, str] = {}
                    for idx, header in enumerate(headers):
                        row_obj[header] = raw_row[idx] if idx < len(raw_row) else ""
                    rows.append(row_obj)
            tables.append({"headers": headers_raw, "rows": rows, "raw": table_data})

        text = _truncate_text(_normalize_newlines("\n".join(paragraphs).strip()), max_text_chars)
        records = _records_from_doc_tables(tables, max_rows=max_rows)
        if not records:
            records = _records_from_text_key_value(text, max_rows=max_rows)
        return file_kind, _unified_extracted_payload(file_kind=file_kind, text=text, tables=tables, records=records)

    if file_kind in {"png", "jpg", "tiff"}:
        from app.services.sber_extract import extract_image_table_rows_best_effort, ocr_text_from_image_variants

        try:
            table_rows = extract_image_table_rows_best_effort(contents)
        except Exception:
            table_rows = []
        if table_rows:
            capped = table_rows if max_rows is None else table_rows[:max_rows]
            records = [{str(k): str(v) for k, v in r.items()} for r in capped]
            preview = "\n".join(" | ".join(r.values()) for r in records[:5])
            text = _truncate_text(_normalize_newlines(preview), max_text_chars)
            return file_kind, _unified_extracted_payload(
                file_kind=file_kind,
                text=text,
                records=records,
                extra_metadata={"image_extraction": "img2table"},
            )

        text_sber = ""
        try:
            text_sber = ocr_text_from_image_variants(contents)
        except Exception:
            text_sber = ""
        if (text_sber or "").strip():
            text = _truncate_text(_normalize_newlines(text_sber), max_text_chars)
            records = _records_from_ocr_text(text, max_rows=max_rows)
            return file_kind, _unified_extracted_payload(
                file_kind=file_kind,
                text=text,
                records=records,
                extra_metadata={"image_extraction": "sber_ocr_variants"},
            )

        text = _extract_image_transcript(contents, file_kind)
        text = _truncate_text(_normalize_newlines(text), max_text_chars)
        records = _records_from_ocr_text(text, max_rows=max_rows)
        return file_kind, _unified_extracted_payload(file_kind=file_kind, text=text, records=records)

    if file_kind == "unknown":
        try:
            text = _normalize_newlines(_decode_text_contents(contents))
        except Exception:
            text = ""
        text = _truncate_text(text, max_text_chars)
        if not (text or "").strip():
            raise ParseFileError(
                code="UNSUPPORTED_FILE_TYPE",
                message=f"Unsupported file type. Supported: {', '.join(SUPPORTED_FILE_KINDS)}",
            )
        records = _records_from_text_key_value(text, max_rows=max_rows)
        return "unknown", _unified_extracted_payload(
            file_kind="unknown",
            text=text,
            records=records,
            warnings=["unknown_extension_decoded_as_text"],
        )

    raise ParseFileError(
        code="UNSUPPORTED_FILE_TYPE",
        message=f"Unsupported file type. Supported: {', '.join(SUPPORTED_FILE_KINDS)}",
    )


async def extract_extracted_input(
    file: UploadFile,
    *,
    max_rows: Optional[int] = None,
    max_text_chars: Optional[int] = None,
) -> Tuple[str, Any]:
    contents = await file.read()
    return extract_extracted_input_from_bytes(
        file.filename,
        file.content_type,
        contents,
        max_rows=max_rows,
        max_text_chars=max_text_chars,
    )

